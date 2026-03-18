from __future__ import annotations

import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.draft_structure import parse_markdown_to_blocks
from app.core.models import DraftBlock, DraftDocument, DraftSection


def build_markdown(draft: DraftDocument) -> str:
    lines: list[str] = []
    title = draft.project_name or "投标文件初稿"
    lines.append(f"# {title}")
    lines.append("")
    lines.append("## 投标文件初稿")
    lines.append("")
    lines.append(f"- 项目名称：{draft.project_name or '待补充'}")
    lines.append(f"- 项目编号：{draft.project_number or '待补充'}")
    lines.append(f"- 采购人：{draft.purchaser or '待补充'}")
    lines.append(f"- 代理机构：{draft.agency or '待补充'}")
    lines.append(f"- 投标人：{draft.bidder_name or '待补充'}")
    lines.append(f"- 生成时间：{draft.generated_at or '待补充'}")
    lines.append("")
    if draft.notes:
        lines.append("## 编制说明")
        lines.append("")
        for note in draft.notes:
            lines.append(f"- {note}")
        lines.append("")

    lines.append("## 目录")
    lines.append("")
    for section in draft.sections:
        lines.append(f"1. {section.display_title}")
    lines.append("")

    for section in draft.sections:
        lines.append(f"# {section.display_title}")
        lines.append("")
        if section.content_md.strip():
            lines.append(section.content_md.strip())
        else:
            lines.append("（本章节暂未生成内容）")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def export_docx(draft: DraftDocument, output_path: str) -> None:
    doc = Document()
    doc.core_properties.title = draft.project_name or ""
    _setup_document(doc)
    _add_cover(doc, draft)
    _add_toc(doc, draft.sections)
    for idx, section in enumerate(draft.sections, start=1):
        doc.add_page_break()
        heading = doc.add_heading(section.display_title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _render_blocks(doc, idx, section)
    doc.save(output_path)


def _setup_document(doc: Document) -> None:
    _set_page_layout(doc)
    _configure_styles(doc)
    _add_footer_page_number(doc)


def _set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    section.start_type = WD_SECTION.NEW_PAGE
    section.different_first_page_header_footer = True


def _configure_styles(doc: Document) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.first_line_indent = Cm(0.74)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "黑体"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.first_line_indent = Cm(0)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "黑体"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.first_line_indent = Cm(0)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "楷体"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.paragraph_format.space_before = Pt(6)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.paragraph_format.first_line_indent = Cm(0)

    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.name = "宋体"
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    list_bullet.font.size = Pt(12)
    list_bullet.paragraph_format.line_spacing = 1.5
    list_bullet.paragraph_format.first_line_indent = Cm(0)

    list_number = doc.styles["List Number"]
    list_number.font.name = "宋体"
    list_number._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    list_number.font.size = Pt(12)
    list_number.paragraph_format.line_spacing = 1.5
    list_number.paragraph_format.first_line_indent = Cm(0)


def _add_cover(doc: Document, draft: DraftDocument) -> None:
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(draft.project_name or "投标文件初稿")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("投标文件初稿")
    run2.bold = True
    run2.font.name = "黑体"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run2.font.size = Pt(18)

    doc.add_paragraph("")
    doc.add_paragraph("")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("项目名称", draft.project_name or "待补充"),
        ("项目编号", draft.project_number or "待补充"),
        ("采购人", draft.purchaser or "待补充"),
        ("代理机构", draft.agency or "待补充"),
        ("投标人", draft.bidder_name or "（盖章）待补充"),
        ("生成时间", draft.generated_at or "待补充"),
    ]
    if len(table.rows) < len(rows):
        for _ in range(len(rows) - len(table.rows)):
            table.add_row()
    for row, (k, v) in zip(table.rows, rows):
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _render_inline_markdown(row.cells[0].paragraphs[0], k, default_bold=True)
        _render_inline_markdown(row.cells[1].paragraphs[0], v)

    doc.add_paragraph("")
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.paragraph_format.first_line_indent = Cm(0)
    sign.add_run("投标人（盖章）：").bold = True
    sign.add_run(draft.bidder_name or "________________")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.first_line_indent = Cm(0)
    date_p.add_run("日期：").bold = True
    date_p.add_run(_display_date(draft.generated_at))

    if draft.notes:
        doc.add_paragraph("")
        h = doc.add_paragraph("编制说明", style="Heading 2")
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for note in draft.notes:
            _add_formatted_paragraph(doc, note, style="List Bullet")

    doc.add_page_break()


def _add_toc(doc: Document, sections: Iterable[DraftSection]) -> None:
    heading = doc.add_heading("目录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, section in enumerate(sections, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        left = p.add_run(f"{idx}. {section.display_title}")
        left.font.name = "宋体"
        left._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _render_blocks(doc: Document, section_index: int, section: DraftSection) -> None:
    blocks = parse_markdown_to_blocks(section.content_md) if section.content_md else section.blocks
    if not blocks:
        _add_formatted_paragraph(doc, "（本章节暂未生成内容）")
        return

    for block in blocks:
        if block.block_type == "heading":
            level = min(max(block.level or 1, 1), 3)
            style = "Heading 2" if level == 1 else "Heading 3"
            title = f"{block.number} {block.text}".strip() if block.number else block.text
            _add_formatted_paragraph(doc, title, style=style)
        elif block.block_type == "paragraph":
            _add_formatted_paragraph(doc, block.text)
        elif block.block_type == "bullet_list":
            for item in block.items:
                _add_formatted_paragraph(doc, item, style="List Bullet")
        elif block.block_type == "number_list":
            for row in block.rows:
                number = row[0] if len(row) > 0 else ""
                title = row[1] if len(row) > 1 else ""
                body = row[2] if len(row) > 2 else ""
                _add_formatted_paragraph(doc, f"{number}. {title}")
                if body:
                    _add_formatted_paragraph(doc, body)
        elif block.block_type == "table":
            _add_table_rows(doc, block.rows)


def _is_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def _add_table_rows(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(rows):
        for col_idx, cell in enumerate(row):
            table_cell = table.cell(row_idx, col_idx)
            table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = table_cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.25
            _render_inline_markdown(paragraph, cell)
            if row_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
                _shade_cell(table_cell, "D9EAF7")


def _add_formatted_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    if style not in {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}:
        paragraph.paragraph_format.line_spacing = 1.5
        if _starts_with_manual_number(text):
            paragraph.paragraph_format.first_line_indent = Cm(0)
    _render_inline_markdown(paragraph, text)


def _render_inline_markdown(paragraph, text: str, default_bold: bool = False) -> None:
    parts = re.split(r"(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)", text)
    for part in parts:
        if not part:
            continue
        bold = False
        italic = False
        value = part
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            bold = True
            value = part[2:-2]
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            italic = True
            value = part[1:-1]
        run = paragraph.add_run(value)
        run.bold = bold or default_bold
        run.italic = italic
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_footer_page_number(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("第 ")
        _add_page_field(paragraph, "PAGE")
        paragraph.add_run(" 页")

        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_p.paragraph_format.first_line_indent = Cm(0)
        run = header_p.add_run(f"{_header_project_name(doc)}  |  投标文件初稿")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def _add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _display_date(text: str) -> str:
    if not text:
        return "待补充"
    return text.split(" ")[0]


def _header_project_name(doc: Document) -> str:
    core = doc.core_properties
    return core.title or "项目名称待补充"


def _starts_with_manual_number(text: str) -> bool:
    stripped = text.strip()
    return bool(
        re.match(r"^\d+(?:\.\d+)*[\.]?\s+", stripped)
        or re.match(r"^[（(]\d+[）)]", stripped)
    )
