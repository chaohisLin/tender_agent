from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple
from docx import Document
from pypdf import PdfReader
from app.core.models import RequirementItem, RequirementType, SourceSpan, OutlineSection, ResponseMatrixRow


class TenderDocumentReader:
    @staticmethod
    def read(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".txt" or suffix == ".md":
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            return TenderDocumentReader._read_docx(path)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        raise ValueError(f"暂不支持的文件类型: {suffix}")

    @staticmethod
    def _read_docx(path: Path) -> str:
        doc = Document(path)
        blocks: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    blocks.append(" | ".join(cells))

        return "\n".join(blocks)


class ParserAgent:
    type_patterns = {
        RequirementType.technical: [r"技术要求", r"性能要求", r"功能要求", r"实施方案"],
        RequirementType.business: [r"商务", r"交付", r"付款", r"工期", r"售后"],
        RequirementType.scoring: [r"评分", r"评审", r"分值"],
        RequirementType.format: [r"格式", r"目录", r"签章", r"附件", r"响应矩阵"],
    }

    def split_paragraphs(self, text: str) -> List[str]:
        items = [p.strip() for p in re.split(r"\n{1,}", text) if p.strip()]
        return [p for p in items if len(p) >= 6]

    def classify(self, text: str) -> Tuple[RequirementType, int, bool]:
        mandatory = bool(re.search(r"必须|不得|应|须|需|不得低于|不得超过", text))
        for req_type, patterns in self.type_patterns.items():
            if any(re.search(p, text) for p in patterns):
                return req_type, (95 if mandatory else 70), mandatory
        return RequirementType.other, (85 if mandatory else 50), mandatory

    def parse_requirements(self, text: str) -> List[RequirementItem]:
        requirements = []
        for i, para in enumerate(self.split_paragraphs(text), start=1):
            req_type, priority, mandatory = self.classify(para)
            requirements.append(
                RequirementItem(
                    req_id=f"REQ-{i:03d}",
                    title=para[:28],
                    content=para,
                    requirement_type=req_type,
                    priority=priority,
                    mandatory=mandatory,
                    source=SourceSpan(paragraph_id=f"P{i}", quote=para[:160]),
                )
            )
        return requirements

    def build_outline(self, requirements: List[RequirementItem]) -> List[OutlineSection]:
        by_type = {}
        for item in requirements:
            by_type.setdefault(item.requirement_type, []).append(item.req_id)
        return [
            OutlineSection(section_id="S1", title="项目理解与总体响应", objective="概述需求理解、实施边界与总体承诺", related_requirements=by_type.get(RequirementType.other, [])[:3]),
            OutlineSection(section_id="S2", title="技术方案", objective="针对技术和功能类要求给出实施方案", related_requirements=by_type.get(RequirementType.technical, [])),
            OutlineSection(section_id="S3", title="商务与实施计划", objective="说明工期、交付、服务与商务响应", related_requirements=by_type.get(RequirementType.business, [])),
            OutlineSection(section_id="S4", title="评分点应答", objective="按评分项强化表达并突出可证明优势", related_requirements=by_type.get(RequirementType.scoring, [])),
            OutlineSection(section_id="S5", title="格式与附件说明", objective="补齐目录、签章、附件、矩阵等要求", related_requirements=by_type.get(RequirementType.format, [])),
        ]

    def build_response_matrix(self, requirements: List[RequirementItem], outline: List[OutlineSection]) -> List[ResponseMatrixRow]:
        rows = []
        section_map = {}
        for s in outline:
            for req_id in s.related_requirements:
                section_map[req_id] = s.title
        for req in requirements:
            rows.append(
                ResponseMatrixRow(
                    req_id=req.req_id,
                    requirement_summary=req.content[:100],
                    response_section=section_map.get(req.req_id, "待补充"),
                    response_text="待生成",
                    evidence_ids=[],
                )
            )
        return rows

    def extract_metadata(self, text: str) -> dict:
        def search(pattern: str) -> str:
            match = re.search(pattern, text, re.MULTILINE)
            return match.group(1).strip() if match else ""

        return {
            "project_name": search(r"项目名称[：:\s]+([^\n]+)"),
            "project_number": search(r"项目编号/包号[：:\s]+([^\n]+)"),
            "purchaser": search(r"采\s*购\s*人[：:\s]+([^\n]+)"),
            "agency": search(r"采购代理机构[：:\s]+([^\n]+)"),
        }
